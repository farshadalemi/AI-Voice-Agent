"""
File processing service for Data Integration Service
Handles parsing and processing of various file formats
"""

import pandas as pd
import json
import PyPDF2
import docx
from typing import Dict, List, Any, Optional, Tuple
import logging
import os
from datetime import datetime
import asyncio
from sqlalchemy.orm import Session

from app.core.database import get_db_session
from app.models.database import DataSource, DataChunk
from app.services.vector_service import VectorService

logger = logging.getLogger(__name__)


class FileProcessor:
    """Service for processing uploaded files"""
    
    def __init__(self):
        self.vector_service = VectorService()
        self.chunk_size = 1000  # Characters per chunk
        self.overlap_size = 200  # Overlap between chunks
    
    async def process_file(self, data_source_id: str, file_path: str, business_id: str) -> bool:
        """Process a file and extract data"""
        
        try:
            # Get database session
            db = next(get_db_session())
            
            # Get data source
            data_source = db.query(DataSource).filter(
                DataSource.id == data_source_id
            ).first()
            
            if not data_source:
                logger.error(f"Data source {data_source_id} not found")
                return False
            
            # Update status to processing
            data_source.processing_status = "processing"
            data_source.metadata["processing_started_at"] = datetime.utcnow().isoformat()
            db.commit()
            
            # Process based on file type
            file_extension = data_source.source_type.lower()
            
            if file_extension in ['xlsx', 'xls']:
                data = await self._process_excel(file_path)
            elif file_extension == 'csv':
                data = await self._process_csv(file_path)
            elif file_extension == 'json':
                data = await self._process_json(file_path)
            elif file_extension == 'pdf':
                data = await self._process_pdf(file_path)
            elif file_extension == 'txt':
                data = await self._process_text(file_path)
            elif file_extension == 'docx':
                data = await self._process_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Create chunks for vector search
            chunks = self._create_chunks(data, data_source_id, business_id)
            
            # Save chunks to database
            for chunk in chunks:
                db_chunk = DataChunk(
                    data_source_id=data_source_id,
                    business_id=business_id,
                    content=chunk["content"],
                    content_hash=chunk["hash"],
                    metadata=chunk["metadata"],
                    chunk_index=chunk["index"]
                )
                db.add(db_chunk)
            
            # Update data source
            data_source.processing_status = "completed"
            data_source.records_count = len(chunks)
            data_source.processed_at = datetime.utcnow()
            data_source.metadata["processing_completed_at"] = datetime.utcnow().isoformat()
            data_source.metadata["chunks_created"] = len(chunks)
            
            db.commit()
            
            # Index in vector database
            await self._index_chunks(chunks, data_source_id)
            
            logger.info(f"Successfully processed file {file_path} with {len(chunks)} chunks")
            return True
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            
            # Update error status
            try:
                data_source.processing_status = "error"
                data_source.processing_error = str(e)
                data_source.metadata["processing_error_at"] = datetime.utcnow().isoformat()
                db.commit()
            except:
                pass
            
            return False
        
        finally:
            db.close()
    
    async def _process_excel(self, file_path: str) -> List[Dict[str, Any]]:
        """Process Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            all_data = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to records
                records = df.to_dict('records')
                
                # Add sheet metadata
                for record in records:
                    record['_sheet_name'] = sheet_name
                    record['_source_type'] = 'excel'
                
                all_data.extend(records)
            
            return all_data
            
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {str(e)}")
            raise
    
    async def _process_csv(self, file_path: str) -> List[Dict[str, Any]]:
        """Process CSV file"""
        try:
            df = pd.read_csv(file_path)
            records = df.to_dict('records')
            
            # Add metadata
            for record in records:
                record['_source_type'] = 'csv'
            
            return records
            
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {str(e)}")
            raise
    
    async def _process_json(self, file_path: str) -> List[Dict[str, Any]]:
        """Process JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Handle different JSON structures
            if isinstance(data, list):
                records = data
            elif isinstance(data, dict):
                records = [data]
            else:
                records = [{"content": str(data)}]
            
            # Add metadata
            for record in records:
                if isinstance(record, dict):
                    record['_source_type'] = 'json'
                else:
                    record = {"content": str(record), "_source_type": "json"}
            
            return records
            
        except Exception as e:
            logger.error(f"Error processing JSON file {file_path}: {str(e)}")
            raise
    
    async def _process_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Process PDF file"""
        try:
            text_content = ""
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
            return [{
                "content": text_content,
                "_source_type": "pdf",
                "_pages": len(pdf_reader.pages)
            }]
            
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            raise
    
    async def _process_text(self, file_path: str) -> List[Dict[str, Any]]:
        """Process text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return [{
                "content": content,
                "_source_type": "text"
            }]
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise
    
    async def _process_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Process Word document"""
        try:
            doc = docx.Document(file_path)
            
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            return [{
                "content": content,
                "_source_type": "docx",
                "_paragraphs": len(doc.paragraphs)
            }]
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise
    
    def _create_chunks(self, data: List[Dict[str, Any]], data_source_id: str, business_id: str) -> List[Dict[str, Any]]:
        """Create text chunks for vector indexing"""
        chunks = []
        
        for i, record in enumerate(data):
            # Convert record to text
            if isinstance(record, dict):
                if 'content' in record:
                    text = record['content']
                else:
                    # Convert dict to readable text
                    text = "\n".join([f"{k}: {v}" for k, v in record.items() if not k.startswith('_')])
            else:
                text = str(record)
            
            # Create chunks from text
            text_chunks = self._split_text(text)
            
            for j, chunk_text in enumerate(text_chunks):
                chunk = {
                    "content": chunk_text,
                    "hash": self._get_content_hash(chunk_text),
                    "index": len(chunks),
                    "metadata": {
                        "record_index": i,
                        "chunk_index": j,
                        "source_type": record.get('_source_type', 'unknown'),
                        "original_record": record
                    }
                }
                chunks.append(chunk)
        
        return chunks
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at word boundary
            if end < len(text):
                # Find last space before end
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start with overlap
            start = end - self.overlap_size if end < len(text) else end
        
        return chunks
    
    def _get_content_hash(self, content: str) -> str:
        """Generate hash for content"""
        import hashlib
        return hashlib.sha256(content.encode()).hexdigest()
    
    async def _index_chunks(self, chunks: List[Dict[str, Any]], data_source_id: str):
        """Index chunks in vector database"""
        try:
            # This would integrate with the vector service
            # For now, we'll just log
            logger.info(f"Would index {len(chunks)} chunks for data source {data_source_id}")
            
            # TODO: Implement actual vector indexing
            # await self.vector_service.index_chunks(chunks, data_source_id)
            
        except Exception as e:
            logger.error(f"Error indexing chunks: {str(e)}")
